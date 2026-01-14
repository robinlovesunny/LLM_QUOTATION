"""
Multimodal Information Extractor
Extracts quotation-related information from various file formats:
- Images (PNG, JPG, JPEG) - via qwen-vl-max
- PDF documents
- Word documents (DOC, DOCX)
- Text files (TXT)
- Excel files (XLS, XLSX)
"""
import os
import io
import base64
import tempfile
from typing import Dict, Any, List, Optional, BinaryIO
from pathlib import Path
from loguru import logger

import dashscope
from dashscope import MultiModalConversation

from app.core.config import settings


# Configure API Key
dashscope.api_key = settings.DASHSCOPE_API_KEY

# Extraction prompt template
EXTRACTION_PROMPT = """Please analyze this document/image and extract quotation-related information.

Extract the following fields if present:
1. Product names/models mentioned
2. Quantities
3. Unit prices
4. Total amounts
5. Customer/company name
6. Contact information
7. Date/validity period
8. Any special terms or discounts

Return the extracted information in a structured JSON format like:
{
    "products": [{"name": "...", "quantity": ..., "unit_price": ..., "total": ...}],
    "customer": {"name": "...", "contact": "...", "email": "..."},
    "quote_date": "...",
    "validity": "...",
    "total_amount": ...,
    "notes": "...",
    "raw_text": "..." // Original text content extracted
}

If certain fields are not found, set them to null.
Respond ONLY with the JSON, no additional text."""


class MultimodalExtractor:
    """Multimodal information extractor using Qwen VL models"""
    
    SUPPORTED_IMAGE_TYPES = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}
    SUPPORTED_DOC_TYPES = {'.pdf', '.doc', '.docx', '.txt', '.xls', '.xlsx', '.csv'}
    
    def __init__(self):
        self.vl_model = "qwen-vl-max"
        self.text_model = "qwen-max"
    
    async def extract_from_file(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract information from uploaded file
        
        Args:
            file_content: File binary content
            filename: Original filename
            mime_type: MIME type of the file
            
        Returns:
            Extracted information dict
        """
        file_ext = Path(filename).suffix.lower()
        logger.info(f"[Extractor] Processing file: {filename}, type: {file_ext}")
        
        try:
            if file_ext in self.SUPPORTED_IMAGE_TYPES:
                return await self._extract_from_image(file_content, filename)
            elif file_ext == '.pdf':
                return await self._extract_from_pdf(file_content, filename)
            elif file_ext in {'.doc', '.docx'}:
                return await self._extract_from_word(file_content, filename)
            elif file_ext == '.txt':
                return await self._extract_from_text(file_content, filename)
            elif file_ext in {'.xls', '.xlsx', '.csv'}:
                return await self._extract_from_excel(file_content, filename)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_ext}",
                    "supported_types": list(self.SUPPORTED_IMAGE_TYPES | self.SUPPORTED_DOC_TYPES)
                }
        except Exception as e:
            logger.error(f"[Extractor] Error processing {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }
    
    async def _extract_from_image(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Extract information from image using Qwen-VL"""
        logger.info(f"[Extractor] Using {self.vl_model} for image extraction")
        
        # Encode image to base64
        base64_image = base64.b64encode(content).decode('utf-8')
        
        # Determine image format
        file_ext = Path(filename).suffix.lower().lstrip('.')
        if file_ext == 'jpg':
            file_ext = 'jpeg'
        
        # Build multimodal message
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "image": f"data:image/{file_ext};base64,{base64_image}"
                    },
                    {
                        "text": EXTRACTION_PROMPT
                    }
                ]
            }
        ]
        
        try:
            response = MultiModalConversation.call(
                model=self.vl_model,
                messages=messages
            )
            
            if response.status_code == 200:
                content_text = response.output.choices[0].message.content[0].get("text", "")
                extracted = self._parse_extraction_result(content_text)
                return {
                    "success": True,
                    "source_type": "image",
                    "filename": filename,
                    "extracted_data": extracted,
                    "raw_response": content_text
                }
            else:
                return {
                    "success": False,
                    "error": f"API error: {response.message}",
                    "filename": filename
                }
        except Exception as e:
            logger.error(f"[Extractor] Image extraction failed: {e}")
            return {"success": False, "error": str(e), "filename": filename}
    
    async def _extract_from_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Extract information from PDF"""
        logger.info("[Extractor] Extracting from PDF")
        
        try:
            import pypdf
            
            # Read PDF from bytes
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            text_content = []
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text() or "")
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                return {
                    "success": False,
                    "error": "PDF appears to be image-based. Please upload as image.",
                    "filename": filename
                }
            
            # Use text model to extract structured data
            return await self._extract_from_text_content(full_text, filename, "pdf")
            
        except ImportError:
            return {
                "success": False,
                "error": "PDF processing requires pypdf library. Install with: pip install pypdf",
                "filename": filename
            }
    
    async def _extract_from_word(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Extract information from Word document"""
        logger.info("[Extractor] Extracting from Word document")
        
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            text_content = []
            
            for para in doc.paragraphs:
                text_content.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            return await self._extract_from_text_content(full_text, filename, "docx")
            
        except ImportError:
            return {
                "success": False,
                "error": "Word processing requires python-docx library. Install with: pip install python-docx",
                "filename": filename
            }
    
    async def _extract_from_text(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Extract information from plain text file"""
        logger.info("[Extractor] Extracting from text file")
        
        # Try different encodings
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                text = content.decode(encoding)
                return await self._extract_from_text_content(text, filename, "txt")
            except UnicodeDecodeError:
                continue
        
        return {
            "success": False,
            "error": "Unable to decode text file",
            "filename": filename
        }
    
    async def _extract_from_excel(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Extract information from Excel file"""
        logger.info("[Extractor] Extracting from Excel file")
        
        try:
            import pandas as pd
        except ImportError:
            return {
                "success": False,
                "error": "Excel processing requires pandas. Install with: pip install pandas openpyxl",
                "filename": filename
            }
        
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.csv':
                # Try different encodings for CSV
                df = None
                for encoding in ['utf-8', 'gbk', 'gb2312']:
                    try:
                        df = pd.read_csv(io.BytesIO(content), encoding=encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                if df is None:
                    return {"success": False, "error": "Unable to decode CSV file", "filename": filename}
            else:
                df = pd.read_excel(io.BytesIO(content), engine='openpyxl')
            
            # Convert DataFrame to text representation
            text_content = df.to_string()
            
            # Also include CSV format for better parsing
            csv_content = df.to_csv(index=False)
            
            full_text = f"Table Data:\n{text_content}\n\nCSV Format:\n{csv_content}"
            return await self._extract_from_text_content(full_text, filename, "excel")
            
        except Exception as e:
            logger.error(f"[Extractor] Excel extraction error: {e}")
            return {
                "success": False,
                "error": f"Excel processing error: {str(e)}",
                "filename": filename
            }
    
    async def _extract_from_text_content(
        self, 
        text: str, 
        filename: str, 
        source_type: str
    ) -> Dict[str, Any]:
        """Use text model to extract structured data from text content"""
        from dashscope import Generation
        
        # Truncate if too long
        max_length = 30000
        if len(text) > max_length:
            text = text[:max_length] + "\n...[truncated]"
        
        messages = [
            {"role": "system", "content": "You are an expert at extracting structured information from documents."},
            {"role": "user", "content": f"{EXTRACTION_PROMPT}\n\nDocument content:\n{text}"}
        ]
        
        try:
            response = Generation.call(
                model=self.text_model,
                messages=messages,
                result_format="message"
            )
            
            if response.status_code == 200:
                content_text = response.output.choices[0].message.content
                extracted = self._parse_extraction_result(content_text)
                return {
                    "success": True,
                    "source_type": source_type,
                    "filename": filename,
                    "extracted_data": extracted,
                    "raw_text": text[:5000],  # Include first 5000 chars
                    "raw_response": content_text
                }
            else:
                return {
                    "success": False,
                    "error": f"API error: {response.message}",
                    "filename": filename
                }
        except Exception as e:
            logger.error(f"[Extractor] Text extraction failed: {e}")
            return {"success": False, "error": str(e), "filename": filename}
    
    def _parse_extraction_result(self, response_text: str) -> Dict[str, Any]:
        """Parse JSON from model response"""
        import json
        import re
        
        # Try to extract JSON from response
        try:
            # Try direct parse first
            return json.loads(response_text)
        except json.JSONDecodeError:
            pass
        
        # Try to find JSON in code blocks
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', response_text)
        if json_match:
            try:
                return json.loads(json_match.group(1))
            except json.JSONDecodeError:
                pass
        
        # Try to find any JSON object in response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            try:
                return json.loads(json_match.group(0))
            except json.JSONDecodeError:
                pass
        
        # Return raw text if parsing fails
        return {"raw_text": response_text, "parse_error": True}
    
    @staticmethod
    def get_supported_types() -> Dict[str, List[str]]:
        """Return supported file types"""
        return {
            "images": [".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"],
            "documents": [".pdf", ".doc", ".docx", ".txt"],
            "spreadsheets": [".xls", ".xlsx", ".csv"]
        }


# Global instance
multimodal_extractor = MultimodalExtractor()
